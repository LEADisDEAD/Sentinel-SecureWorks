from fastapi import FastAPI
from .database import engine, Base
from . import models
from .routes import auth_routes, questionnaire_routes
from .routes import reference_routes, answer_routes
from .auth import get_current_user
from fastapi import Depends,Form
from fastapi.templating import Jinja2Templates
from fastapi import Request
from fastapi import Request, Depends
from sqlalchemy.orm import Session
from app.database import SessionLocal
from app.auth import get_current_user
from app import models
from starlette.middleware.sessions import SessionMiddleware
from fastapi.responses import RedirectResponse
from app.models import Questionnaire, ReferenceDocument
from app.routes.answer_routes import compute_confidence

templates = Jinja2Templates(directory="app/templates")

Base.metadata.create_all(bind=engine)

app = FastAPI()

from fastapi.staticfiles import StaticFiles

app.mount("/static", StaticFiles(directory="app/static"), name="static")

app.add_middleware(SessionMiddleware, secret_key="super-secret-key")
app.include_router(auth_routes.router)
app.include_router(questionnaire_routes.router)
app.include_router(reference_routes.router)
app.include_router(answer_routes.router)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

@app.get("/")
def root():
    return RedirectResponse(url="/auth", status_code=303)

@app.get("/protected")
def protected_route(current_user = Depends(get_current_user)):
    return {"message": f"Hello {current_user.email}, you are authenticated"}

# ===========================================================

@app.get("/auth")
def auth_page(request: Request):
    return templates.TemplateResponse(
        "auth.html",
        {
            "request": request,
            "hide_navbar": True
        }
    )

@app.post("/login-ui")
def login_ui_post(
    request: Request,
    email: str = Form(...),
    password: str = Form(...),
    db: Session = Depends(get_db)
):
    user = db.query(models.User).filter(models.User.email == email).first()

    from app.auth import verify_password

    if not user or not verify_password(password, user.hashed_password):
        return templates.TemplateResponse(
            "auth.html",
            {
                "request": request,
                "error": "Invalid credentials",
                "hide_navbar": True
            }
        )

    request.session["user_id"] = user.id
    return RedirectResponse(url="/dashboard", status_code=303)


@app.post("/signup-ui")
def signup_ui_post(
    request: Request,
    email: str = Form(...),
    password: str = Form(...),
    db: Session = Depends(get_db)
):
    from app.auth import hash_password

    existing = db.query(models.User).filter(models.User.email == email).first()
    if existing:
        return templates.TemplateResponse(
            "auth.html",
            {
                "request": request,
                "error": "Invalid credentials",
                "hide_navbar": True
            }
        )

    new_user = models.User(
        email=email,
        hashed_password=hash_password(password)
    )

    db.add(new_user)
    db.commit()
    db.refresh(new_user)

    request.session["user_id"] = new_user.id

    return RedirectResponse(url="/dashboard", status_code=303)

# =========================================================

@app.get("/dashboard")
def dashboard(
    request: Request,
    db: Session = Depends(get_db)
):
    user_id = request.session.get("user_id")

    if not user_id:
        return RedirectResponse(url="/auth", status_code=303)

    user = db.query(models.User).filter(models.User.id == user_id).first()

    questionnaires = db.query(models.Questionnaire).filter(
        models.Questionnaire.user_id == user.id
    ).all()

    dashboard_data = []

    for q in questionnaires:
        total_questions = len(q.questions)
        answered = 0

        for question in q.questions:
            answer = db.query(models.Answer).filter(
                models.Answer.question_id == question.id
            ).order_by(models.Answer.id.desc()).first()

            if answer and answer.answer_text:
                answered += 1

        completion = 0
        if total_questions > 0:
            completion = round((answered / total_questions) * 100, 2)

        dashboard_data.append({
            "id": q.id,
            "title": q.title,
            "completion": completion
        })
    
    if dashboard_data:
        avg_completion = sum(q["completion"] for q in dashboard_data) / len(dashboard_data)
    else:
        avg_completion = 0
        
    references = db.query(models.ReferenceDocument).filter(models.ReferenceDocument.user_id == user.id).all()

    return templates.TemplateResponse(
        "dashboard.html",
        {
            "request": request,
            "user": user,
            "questionnaires": dashboard_data,
            "references": references,
            "avg_completion": round(avg_completion, 1)
        }
    )
    
from fastapi import UploadFile, File
import PyPDF2

@app.post("/questionnaire/upload-ui")
async def upload_questionnaire_ui(
    request: Request,
    title: str = Form(...),
    content: str = Form(None),
    file: UploadFile = File(None),
    db: Session = Depends(get_db)
):
    user_id = request.session.get("user_id")

    if not user_id:
        return RedirectResponse(url="/auth", status_code=303)

    extracted_text = None

    # Case 1: Manual paste
    if content:
        extracted_text = content

    # Case 2: File upload
    elif file:
        allowed_extensions = (".pdf", ".txt")

        if not file.filename.lower().endswith(allowed_extensions):
            return RedirectResponse(url="/dashboard", status_code=303)

        if file.filename.lower().endswith(".txt"):
            extracted_text = (await file.read()).decode("utf-8")

        elif file.filename.lower().endswith(".pdf"):
            reader = PyPDF2.PdfReader(file.file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            extracted_text = text

    questionnaire = models.Questionnaire(
        title=title,
        user_id=user_id
    )

    db.add(questionnaire)
    db.commit()
    db.refresh(questionnaire)

    questions = [q.strip() for q in extracted_text.split("\n") if q.strip()]

    for q_text in questions:
        question = models.Question(
            text=q_text,
            questionnaire_id=questionnaire.id
        )
        db.add(question)

    db.commit()

    return RedirectResponse(url="/dashboard", status_code=303)

from fastapi import UploadFile, File
import PyPDF2

@app.post("/reference/upload-ui")
async def upload_reference_ui(
    request: Request,
    title: str = Form(None),
    content: str = Form(None),
    file: UploadFile = File(None),
    db: Session = Depends(get_db)
    ):
    user_id = request.session.get("user_id")

    if not user_id:
        return RedirectResponse(url="/auth", status_code=303)

    extracted_text = None

    # Case 1: Manual text
    if content:
        extracted_text = content

    # Case 2: File upload
    elif file:
        if file.filename.endswith(".txt"):
            extracted_text = (await file.read()).decode("utf-8")

        elif file.filename.endswith(".pdf"):
            reader = PyPDF2.PdfReader(file.file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            extracted_text = text
        else:
            return RedirectResponse(url="/dashboard", status_code=303)

    else:
        return RedirectResponse(url="/dashboard", status_code=303)

    reference = models.ReferenceDocument(
        title=title if title else file.filename if file else "Reference Document",
        content=extracted_text,
        user_id=user_id
    )

    db.add(reference)
    db.commit()

    return RedirectResponse(url="/dashboard", status_code=303)

@app.get("/questionnaire/{questionnaire_id}/generate-ui")
def generate_ui(
    questionnaire_id: int,
    request: Request,
    db: Session = Depends(get_db)
 ):
    user_id = request.session.get("user_id")

    if not user_id:
        return RedirectResponse(url="/auth", status_code=303)

    questionnaire = db.query(models.Questionnaire).filter(
        models.Questionnaire.id == questionnaire_id,
        models.Questionnaire.user_id == user_id
    ).first()

    if not questionnaire:
        return RedirectResponse(url="/dashboard", status_code=303)

    references = db.query(models.ReferenceDocument).filter(
        models.ReferenceDocument.user_id == user_id
    ).all()

    if not references:
        return templates.TemplateResponse(
            "dashboard.html",
            {
                "request": request,
                "user": db.query(models.User).filter(models.User.id == user_id).first(),
                "questionnaires": [],
                "references": [],
                "avg_completion": 0,
                "error": "Upload at least one reference document before generating answers."
            }
        )

    from app.utils.groq_client import generate_answer

    reference_text = "\n\n".join([r.content for r in references])
    
    # for rate limiting
    import time

    # Remove old answers before regenerating
    for question in questionnaire.questions:
        db.query(models.Answer).filter(
            models.Answer.question_id == question.id
        ).delete()


    for question in questionnaire.questions:
        response = generate_answer(question.text, reference_text)
        time.sleep(2)
        import json
        try:
            parsed = json.loads(response)
            answer_text = parsed.get("answer")
            citation_text = parsed.get("citation")
        except:
            answer_text = response
            citation_text = None

        answer = models.Answer(
            answer_text=answer_text,
            citation=citation_text,
            question_id=question.id
        )

        db.add(answer)

    db.commit()

    return RedirectResponse(url="/dashboard", status_code=303)


from app.routes.answer_routes import compute_confidence

@app.get("/questionnaire/{questionnaire_id}")
def review_questionnaire(
    questionnaire_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    user_id = request.session.get("user_id")

    if not user_id:
        return RedirectResponse(url="/auth", status_code=303)

    questionnaire = db.query(models.Questionnaire).filter(
        models.Questionnaire.id == questionnaire_id,
        models.Questionnaire.user_id == user_id
    ).first()

    if not questionnaire:
        return RedirectResponse(url="/dashboard", status_code=303)

    total_questions = len(questionnaire.questions)
    answered_with_citation = 0
    answered_without_citation = 0
    not_found = 0

    enriched_questions = []

    for question in questionnaire.questions:
        answer = db.query(models.Answer).filter(
            models.Answer.question_id == question.id
        ).order_by(models.Answer.id.desc()).first()

        confidence = 0.0

        if answer and answer.answer_text:
            confidence = compute_confidence(answer.answer_text, answer.citation)

            if answer.answer_text.strip().lower() == "not found in references.":
                not_found += 1
            elif answer.citation:
                answered_with_citation += 1
            else:
                answered_without_citation += 1

        enriched_questions.append({
            "question": question,
            "answer": answer,
            "confidence": confidence
        })

    completion_percentage = 0
    if total_questions > 0:
        completion_percentage = round(
            ((answered_with_citation + answered_without_citation) / total_questions) * 100,
            2
        )

    summary = {
        "total_questions": total_questions,
        "answered_with_citation": answered_with_citation,
        "answered_without_citation": answered_without_citation,
        "not_found": not_found,
        "completion_percentage": completion_percentage
    }

    return templates.TemplateResponse(
        "review.html",
        {
            "request": request,
            "questionnaire": questionnaire,
            "questions_data": enriched_questions,
            "summary": summary
        }
    )


@app.get("/questionnaire/{questionnaire_id}/export-ui")
def export_ui(
    questionnaire_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    user_id = request.session.get("user_id")

    if not user_id:
        return RedirectResponse(url="/auth", status_code=303)

    questionnaire = db.query(models.Questionnaire).filter(
        models.Questionnaire.id == questionnaire_id,
        models.Questionnaire.user_id == user_id
    ).first()

    if not questionnaire:
        return RedirectResponse(url="/dashboard", status_code=303)

    from docx import Document
    from fastapi.responses import FileResponse

    document = Document()
    document.add_heading(questionnaire.title, level=1)

    for question in questionnaire.questions:
        document.add_heading("Question", level=3)
        document.add_paragraph(question.text)

        # Get latest answer properly
        answer = (
            db.query(models.Answer)
            .filter(models.Answer.question_id == question.id)
            .order_by(models.Answer.id.desc())
            .first()
        )

        if answer and answer.answer_text:
            confidence = compute_confidence(
                answer.answer_text,
                answer.citation or ""
            )

            document.add_heading("Answer", level=4)
            document.add_paragraph(answer.answer_text)

            document.add_paragraph(
                f"Citation: {answer.citation or 'N/A'}"
            )

            document.add_paragraph(
                f"Confidence Score: {confidence}"
            )
        else:
            document.add_paragraph("Not answered.")

        document.add_paragraph("-" * 50)

    import os

    export_dir = "exports"
    os.makedirs(export_dir, exist_ok=True)

    file_path = os.path.join(export_dir, f"export_{questionnaire_id}.docx")
    document.save(file_path)

    return FileResponse(
        path=file_path,
        filename=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.get("/reference/{reference_id}/delete")
def delete_reference(reference_id: int, request: Request, db: Session = Depends(get_db)):
    user_id = request.session.get("user_id")

    ref = db.query(models.ReferenceDocument).filter(
        models.ReferenceDocument.id == reference_id,
        models.ReferenceDocument.user_id == user_id
    ).first()

    if ref:
        db.delete(ref)
        db.commit()

    return RedirectResponse(url="/dashboard", status_code=303)


from fastapi import Depends
from sqlalchemy.orm import Session
from starlette.responses import RedirectResponse

@app.post("/clear-all")
def clear_all(
    request: Request,
    db: Session = Depends(get_db)
 ):
    user_id = request.session.get("user_id")

    if not user_id:
        return RedirectResponse("/auth", status_code=303)

    # Delete questionnaires
    db.query(Questionnaire).filter(
        Questionnaire.user_id == user_id
    ).delete()

    # Delete references
    db.query(ReferenceDocument).filter(
        ReferenceDocument.user_id == user_id
    ).delete()

    db.commit()

    return RedirectResponse("/dashboard", status_code=303)


@app.post("/answers/{answer_id}/edit-ui")
def edit_answer_ui(
    answer_id: int,
    request: Request,
    answer_text: str = Form(...),
    citation: str = Form(None),
    db: Session = Depends(get_db)
):
    user_id = request.session.get("user_id")

    if not user_id:
        return RedirectResponse(url="/auth", status_code=303)

    answer = db.query(models.Answer).filter(
        models.Answer.id == answer_id
    ).first()

    if not answer:
        return RedirectResponse(url="/dashboard", status_code=303)

    questionnaire = answer.question.questionnaire

    if questionnaire.user_id != user_id:
        return RedirectResponse(url="/dashboard", status_code=303)

    answer.answer_text = answer_text
    answer.citation = citation

    db.commit()

    return RedirectResponse(
        url=f"/questionnaire/{questionnaire.id}",
        status_code=303
    )

@app.get("/logout")
def logout(request: Request):
    request.session.clear()
    return RedirectResponse(url="/auth", status_code=303)