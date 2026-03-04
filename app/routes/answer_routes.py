from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from ..database import SessionLocal
from .. import models
from ..auth import get_current_user
from ..utils.groq_client import generate_answer
import json
from app.schemas import AnswerUpdate
import re
from fastapi.responses import FileResponse
from docx import Document
import os


router = APIRouter()


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

#confidence scoring logic

import re

def compute_confidence(answer_text,citation_text):
    if not answer_text:
        return 0.0

    score = 0.3  # base if answer exists

    if citation_text:
        # citation length factor
        citation_length = len(citation_text.split())
        score += min(citation_length / 50, 0.3)  # max +0.3

        # keyword overlap factor
        answer_words = set(re.findall(r'\w+', answer_text.lower()))
        citation_words = set(re.findall(r'\w+', citation_text.lower()))
        overlap = answer_words.intersection(citation_words)

        if len(answer_words) > 0:
            overlap_ratio = len(overlap) / len(answer_words)
            score += min(overlap_ratio, 0.4)  # max +0.4

    return round(min(score, 1.0), 2)


@router.post("/questionnaire/{questionnaire_id}/generate")
def generate_answers(
    questionnaire_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user)
):

    questionnaire = db.query(models.Questionnaire).filter(
        models.Questionnaire.id == questionnaire_id,
        models.Questionnaire.user_id == current_user.id
    ).first()

    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found")

    references = db.query(models.ReferenceDocument).filter(
        models.ReferenceDocument.user_id == current_user.id
    ).all()

    if not references:
        raise HTTPException(status_code=400, detail="No reference documents uploaded")

    reference_text = "\n\n".join([r.content for r in references])

    results = []

    for question in questionnaire.questions:

        try:
            response = generate_answer(question.text, reference_text)
        except Exception as e:
            return {
                "message": "LLM generation failed",
                "error": str(e)
            }

        # Parse model JSON safely
        try:
            parsed = json.loads(response)
            answer_text = parsed.get("answer")
            citation_text = parsed.get("citation")
        except Exception:
            # fallback if model returns plain text
            answer_text = response
            citation_text = None

        answer = models.Answer(
            answer_text=answer_text,
            citation=citation_text,
            question_id=question.id
        )

        db.add(answer)

        # Converting string llm output to proper json object

        try:
            parsed_response = json.loads(response)
        except:
            parsed_response = {
                "raw_response": response
            }

        results.append({
            "question": question.text,
            "llm_response": parsed_response
        })

    db.commit()

    return {
        "message": "Answers generated successfully",
        "results": results
    }
    
@router.get("/questionnaire/{questionnaire_id}/answers")
def get_answers(
    questionnaire_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user)
):
    questionnaire = db.query(models.Questionnaire).filter(
        models.Questionnaire.id == questionnaire_id,
        models.Questionnaire.user_id == current_user.id
    ).first()

    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found")

    results = []
    
    #counters
    total_questions = len(questionnaire.questions)
    answered_with_citation = 0
    answered_without_citation = 0
    not_found = 0

    for question in questionnaire.questions:
    # Get latest answer for question
        answer = db.query(models.Answer).filter(
            models.Answer.question_id == question.id
        ).order_by(models.Answer.id.desc()).first()

        answer_text = None
        citation_text = None
        confidence = 0.0

        if answer and answer.answer_text:
            answer_text = answer.answer_text
            citation_text = answer.citation

            if answer_text.strip().lower() == "not found in references.":
                not_found += 1
            elif citation_text:
                answered_with_citation += 1
            else:
                answered_without_citation += 1

            confidence = compute_confidence(answer_text, citation_text)

        results.append({
            "question_id": question.id,
            "question": question.text,
            "answer": answer_text,
            "citation": citation_text,
            "confidence": confidence
        })
        
    completion_percentage = 0
    if total_questions > 0:
        completion_percentage = round(
            ((answered_with_citation + answered_without_citation) / total_questions) * 100,
            2
        )

    return {
    "questionnaire_id": questionnaire.id,
    "title": questionnaire.title,
    "summary": {
        "total_questions": total_questions,
        "answered_with_citation": answered_with_citation,
        "answered_without_citation": answered_without_citation,
        "not_found": not_found,
        "completion_percentage": completion_percentage
    },
    "results": results
}


@router.put("/answers/{answer_id}")
def update_answer(
    answer_id: int,
    update_data: AnswerUpdate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    answer = db.query(models.Answer).filter(models.Answer.id == answer_id).first()

    if not answer:
        raise HTTPException(status_code=404, detail="Answer not found")

    # Ensure user owns this answer (via questionnaire)
    questionnaire = answer.question.questionnaire
    if questionnaire.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")

    answer.answer_text = update_data.answer
    answer.citation = update_data.citation

    db.commit()
    db.refresh(answer)

    return {
        "message": "Answer updated successfully",
        "answer_id": answer.id,
        "answer": answer.answer_text,
        "citation": answer.citation
    }


@router.get("/questionnaire/{questionnaire_id}/export")
def export_questionnaire(
    questionnaire_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user)
):
    questionnaire = db.query(models.Questionnaire).filter(
        models.Questionnaire.id == questionnaire_id,
        models.Questionnaire.user_id == current_user.id
    ).first()

    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found")

    document = Document()

    document.add_heading(questionnaire.title, level=1)

    for question in questionnaire.questions:
        answer = db.query(models.Answer).filter(
            models.Answer.question_id == question.id
        ).order_by(models.Answer.id.desc()).first()

        document.add_paragraph(f"Question: {question.text}")
        
        if answer:
            document.add_paragraph(f"Answer: {answer.answer_text}")
            if answer.citation:
                document.add_paragraph(f"Citation: {answer.citation}")
        else:
            document.add_paragraph("Answer: Not generated.")

        document.add_paragraph("")  # spacing

    file_path = f"export_questionnaire_{questionnaire_id}.docx"
    document.save(file_path)

    return FileResponse(
        path=file_path,
        filename=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )